from api.users.models import School
from datetime import datetime
from django.conf import settings
from io import BytesIO
from PIL import Image
from typing import Optional, Dict, Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Cm
from docx.text.paragraph import Paragraph

import base64
import os


class InvoiceDocxGenerator:
    IMAGE_WIDTH = 1.5  # cm

    def __init__(self, data: Dict[str, Any], school: School = None):
        """
        Initialize the generator with raw setting data.
        Expects data to be a JSON with possible 'texts', 'tables', and 'images' fields.
        """
        def get_value(key: str):
            return data.get(key, None)

        # Fallback values for sample invoice
        SCHOOL_NAME = "Sample School Name" if not school else school.name
        SCHOOL_ADDRESS = "Sample Address1\nAddress2" if not school else school.address
        STUDENT_COUNT = 1 if not school else school.student_count()

        NOW = datetime.now()
        FEES = get_value("fees")
        TOTAL_FEES = (int(FEES) if FEES else 0) * STUDENT_COUNT

        # Prepare text replacements
        self.texts = {
            "year": str(NOW.year),
            "date": NOW.strftime("%d %B %Y"),
            "school_name": SCHOOL_NAME,
            "school_address": SCHOOL_ADDRESS,
            "total_count": STUDENT_COUNT,
            "total_fees": TOTAL_FEES,
            "account_name": get_value("accountName"),
            "bsb": get_value("bsb"),
            "account_number": get_value("accountNumber"),
        }
        # Prepare table and image replacements
        self.tables = {
            "address": get_value("address"),
            "email": get_value("email"),
            "website": get_value("website"),
            "chair_name": get_value("chairName"),
            "chair_title": get_value("chairTitle"),
        }
        self.images = {
            "signature": get_value("signature")
        }
        self.school_name = SCHOOL_NAME

    def generate_invoice_docx(self) -> BytesIO:
        """
        Generate an invoice DOCX document by replacing placeholders in a Word template
        with provided text, table, and image content.

        Args:
            texts (dict, optional): Dictionary of text replacements in the format {key: value}.
                                    Keys will be converted to placeholders like {{ key }}.
            tables (dict, optional): Dictionary of text replacements for table cells.
            images (dict, optional): Dictionary of image data {key: base64 string or file path}.

        Returns:
            BytesIO: An in-memory DOCX file with all replacements applied.
        """
        template_path = os.path.join(settings.BASE_DIR, "static", "invoice_template.docx")
        doc = Document(template_path) if os.path.exists(template_path) else Document()

        try:
            replacements_text = self._to_placeholder_keys(self.texts)
            replacements_table = self._to_placeholder_keys(self.tables)
            replacements_image = self._to_placeholder_keys(self.images)

            self._replace_in_paragraphs(doc.paragraphs, replacements_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._replace_in_paragraphs(cell.paragraphs, replacements_table)

            self._replace_images(doc, replacements_image)

        except Exception as e:
            print(f"Error generating invoice DOCX: {str(e)}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _replace_in_paragraphs(
        self,
        paragraphs: list[Paragraph],
        replacements: Dict[str, Any],
        clear_format: bool = False
    ):
        """
        Replace placeholders in Word paragraphs.

        This function iterates over each paragraph and replaces specified placeholder strings
        with provided values. Supports both format-preserving and format-clearing replacement.

        Args:
            paragraphs (List[Paragraph]): List of paragraph objects from a docx Document or table cell.
            replacements (Dict[str, Any]): Mapping of placeholders (e.g., '{{ key }}') to values.
            clear_format (bool): If True, clears all formatting and replaces the entire paragraph text.
                                If False, attempts to replace placeholders while preserving formatting.
        """
        for paragraph in paragraphs:
            if not paragraph.text.strip():
                continue

            for key, value in replacements.items():
                if key in paragraph.text:
                    if clear_format:
                        new_text = paragraph.text.replace(key, str(value))
                        for run in paragraph.runs[::-1]:
                            run._element.getparent().remove(run._element)
                        paragraph.add_run(new_text)
                    else:
                        self._replace_text_preserve_format(paragraph, key, str(value))

    def _replace_images(self, doc: DocumentObject, replacements: Dict[str, str]):
        """
        Replace image placeholders in the document with actual images from base64 or file path.

        Args:
            doc (DocumentObject): The Word document object.
            replacements (Dict[str, str]): Mapping of placeholder keys to base64 strings or file paths.
        """
        for placeholder, image_data in replacements.items():
            if not image_data:
                continue

            def process_paragraphs(paragraphs: list[Paragraph]):
                for paragraph in paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, "")
                        self._insert_image(paragraph, image_data, width=self.IMAGE_WIDTH)

            process_paragraphs(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)

    def _insert_image(self, paragraph: Paragraph, image_data: str, width: Optional[float] = None):
        """
        Insert an image into a Word paragraph from base64 string using centimeters.

        Args:
            paragraph (Paragraph): Paragraph to insert image into.
            image_data (str): Base64-encoded image string (with or without data URI prefix).
            width (float, optional): Image width in centimeters. If None, uses original image width.
        """
        try:
            if image_data.startswith("data:image"):
                image_data = image_data.split(",")[1]

            image_bytes = base64.b64decode(image_data)
            image_stream = BytesIO(image_bytes)

            if width is None:
                with Image.open(BytesIO(image_bytes)) as img:
                    dpi = img.info.get("dpi", (96, 96))
                    width_inch = img.width / dpi[0]
                    width = width_inch * 2.54  # Convert inches to cm

            paragraph.add_run().add_picture(image_stream, width=Cm(width))
        except Exception as e:
            print(f"Error adding image: {str(e)}")

    def _replace_text_clear_format(self, paragraph: Paragraph, placeholder: str, replacement: str):
        """
        Replace all instances of a placeholder in a paragraph by wiping all text and formatting.

        Args:
            paragraph (Paragraph): The Word paragraph to modify.
            placeholder (str): Placeholder string to replace.
            replacement (str): The replacement value to insert.
        """
        new_text = paragraph.text.replace(placeholder, replacement)
        for run in paragraph.runs[::-1]:
            run._element.getparent().remove(run._element)
        paragraph.add_run(new_text)

    def _replace_text_preserve_format(self, paragraph: Paragraph, placeholder: str, replacement: str):
        """
        Replace placeholder text in a paragraph while preserving formatting.

        This handles cases where a placeholder may span multiple runs by identifying the runs
        and updating only the relevant sections, preserving all styling where possible.

        Args:
            paragraph (Paragraph): The Word paragraph to modify.
            placeholder (str): Placeholder string to search for and replace.
            replacement (str): The replacement text to insert.
        """
        text_index = paragraph.text.find(placeholder)
        if text_index == -1:
            return

        current_pos = 0
        start_run, end_run = None, None
        offset_start, offset_end = 0, 0

        for i, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            if start_run is None and current_pos <= text_index < current_pos + run_len:
                start_run = i
                offset_start = text_index - current_pos
            if end_run is None and current_pos < text_index + len(placeholder) <= current_pos + run_len:
                end_run = i
                offset_end = (text_index + len(placeholder)) - current_pos
            current_pos += run_len
            if start_run is not None and end_run is not None:
                break

        if start_run is None or end_run is None:
            return

        if start_run == end_run:
            run = paragraph.runs[start_run]
            run.text = run.text[:offset_start] + replacement + run.text[offset_end:]
        else:
            paragraph.runs[start_run].text = paragraph.runs[start_run].text[:offset_start] + replacement
            for i in range(start_run + 1, end_run):
                paragraph.runs[i].text = ""
            paragraph.runs[end_run].text = paragraph.runs[end_run].text[offset_end:]

    def _to_placeholder_keys(self, data: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Convert dictionary keys to the placeholder format used in DOCX templates (e.g., "{{ key }}").

        Args:
            data (Dict[str, Any] or None): Original data dictionary.

        Returns:
            Dict[str, Any]: Dictionary with keys converted to placeholder format.
        """
        return {f"{{{{ {k} }}}}": v for k, v in (data or {}).items()}
